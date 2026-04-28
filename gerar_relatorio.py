from docx import Document
from io import BytesIO


def gerar_relatorio(df_cliente):

    document = Document()

    document.add_heading("Histórico de Cargas - Cliente", level=1)

    num_colunas = len(df_cliente.columns)

    tabela = document.add_table(rows=1, cols=num_colunas)

    celulas_cabecalho = tabela.rows[0].cells

    for i, nome_coluna in enumerate(df_cliente.columns):
        celulas_cabecalho[i].text = nome_coluna


    for _, linha in df_cliente.iterrows():

        nova_linha = tabela.add_row().cells

        for i, valor in enumerate(linha):
            nova_linha[i].text = str(valor)


    # cria arquivo em memória
    buffer = BytesIO()

    # salva no buffer
    document.save(buffer)

    # volta para início
    buffer.seek(0)

    # retorna o arquivo
    return buffer